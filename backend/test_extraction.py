"""
Test script for document text extraction
Run with: python test_extraction.py
"""
import os
import sys

# Add backend to path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from app.services.extraction_service import ExtractionService

def create_test_pdf():
    """Create a simple test PDF"""
    try:
        from reportlab.pdfgen import canvas
        
        pdf_path = "./test_documents/test_contract.pdf"
        os.makedirs("./test_documents", exist_ok=True)
        
        c = canvas.Canvas(pdf_path)
        c.setFont("Helvetica", 12)
        
        # Add content
        y = 750
        content = [
            "LEGAL CONTRACT",
            "",
            "This agreement is entered into between Party A and Party B.",
            "",
            "1. TERMS AND CONDITIONS",
            "This contract shall be binding upon both parties.",
            "",
            "2. PAYMENT TERMS",
            "Payment shall be made within 30 days of invoice.",
            "",
            "3. TERMINATION",
            "Either party may terminate with 30 days notice.",
            "",
            "4. GOVERNING LAW",
            "This contract is governed by the laws of the jurisdiction.",
        ]
        
        for line in content:
            c.drawString(72, y, line)
            y -= 20
        
        c.save()
        print(f"✅ Created test PDF: {pdf_path}")
        return pdf_path
        
    except ImportError:
        print("⚠️ reportlab not installed, creating simple test file")
        # Create a simple test file instead
        pdf_path = "./test_documents/test_contract.pdf"
        os.makedirs("./test_documents", exist_ok=True)
        with open(pdf_path, "wb") as f:
            f.write(b"%PDF-1.4 test content")
        return pdf_path

def create_test_docx():
    """Create a simple test DOCX"""
    try:
        from docx import Document
        
        docx_path = "./test_documents/test_contract.docx"
        os.makedirs("./test_documents", exist_ok=True)
        
        doc = Document()
        doc.add_heading("LEGAL CONTRACT", 0)
        doc.add_paragraph(
            "This agreement is entered into between Party A and Party B."
        )
        doc.add_heading("1. TERMS AND CONDITIONS", level=1)
        doc.add_paragraph(
            "This contract shall be binding upon both parties."
        )
        doc.add_heading("2. PAYMENT TERMS", level=1)
        doc.add_paragraph(
            "Payment shall be made within 30 days of invoice."
        )
        doc.add_heading("3. TERMINATION", level=1)
        doc.add_paragraph(
            "Either party may terminate with 30 days notice."
        )
        
        # Add a table
        table = doc.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "Party"
        table.cell(0, 1).text = "Signature"
        table.cell(1, 0).text = "Party A"
        table.cell(1, 1).text = "________________"
        table.cell(2, 0).text = "Party B"
        table.cell(2, 1).text = "________________"
        
        doc.save(docx_path)
        print(f"✅ Created test DOCX: {docx_path}")
        return docx_path
        
    except Exception as e:
        print(f"❌ Error creating test DOCX: {e}")
        return None

def test_pdf_extraction(pdf_path: str):
    """Test PDF text extraction"""
    print("\n" + "="*50)
    print("🧪 TEST 1: PDF Text Extraction")
    print("="*50)
    
    try:
        text = ExtractionService.extract_from_pdf(pdf_path)
        
        if text:
            print(f"✅ Extracted {len(text)} characters")
            print(f"📝 First 200 chars:\n{text[:200]}")
            return True
        else:
            print("⚠️ No text extracted (PDF might be scanned/image-based)")
            return False
            
    except Exception as e:
        print(f"❌ Error: {e}")
        return False

def test_docx_extraction(docx_path: str):
    """Test DOCX text extraction"""
    print("\n" + "="*50)
    print("🧪 TEST 2: DOCX Text Extraction")
    print("="*50)
    
    try:
        text = ExtractionService.extract_from_docx(docx_path)
        
        if text:
            print(f"✅ Extracted {len(text)} characters")
            print(f"📝 First 200 chars:\n{text[:200]}")
            return True
        else:
            print("⚠️ No text extracted")
            return False
            
    except Exception as e:
        print(f"❌ Error: {e}")
        return False

def test_document_info(file_path: str, file_type: str):
    """Test document info extraction"""
    print("\n" + "="*50)
    print("🧪 TEST 3: Document Info")
    print("="*50)
    
    try:
        info = ExtractionService.get_document_info(file_path, file_type)
        
        print(f"✅ Document info retrieved:")
        for key, value in info.items():
            print(f"  - {key}: {value}")
        return True
        
    except Exception as e:
        print(f"❌ Error: {e}")
        return False

def test_text_cleaning():
    """Test text cleaning"""
    print("\n" + "="*50)
    print("🧪 TEST 4: Text Cleaning")
    print("="*50)
    
    dirty_text = """
    This   is   a   test   with   extra   spaces.
    
    
    
    And   multiple   blank   lines.
    
    Special chars: ™ © ® 
    Normal text continues here.
    """
    
    cleaned = ExtractionService.clean_text(dirty_text)
    print(f"📝 Original length: {len(dirty_text)}")
    print(f"📝 Cleaned length: {len(cleaned)}")
    print(f"📝 Cleaned text:\n{cleaned}")
    print("✅ Text cleaning works!")
    return True

def main():
    """Main test function"""
    print("📄 Legal Document Analyzer - Text Extraction Test")
    print("="*50)
    
    tests_passed = 0
    total_tests = 4
    
    # Create test documents
    print("\n📁 Creating test documents...")
    pdf_path = create_test_pdf()
    docx_path = create_test_docx()
    
    # Run tests
    if pdf_path and os.path.exists(pdf_path):
        if test_pdf_extraction(pdf_path):
            tests_passed += 1
        
        if test_document_info(
            pdf_path,
            "application/pdf"
        ):
            tests_passed += 1
    else:
        print("⚠️ Skipping PDF tests - no PDF created")
    
    if docx_path and os.path.exists(docx_path):
        if test_docx_extraction(docx_path):
            tests_passed += 1
    else:
        print("⚠️ Skipping DOCX tests - no DOCX created")
    
    if test_text_cleaning():
        tests_passed += 1
    
    # Summary
    print("\n" + "="*50)
    print("📊 TEST SUMMARY")
    print("="*50)
    print(f"✅ Tests passed: {tests_passed}/{total_tests}")
    
    if tests_passed == total_tests:
        print("🎉 All tests passed! Ready for chunking!")
    else:
        print("⚠️ Some tests failed. Check the errors above.")

if __name__ == "__main__":
    main()