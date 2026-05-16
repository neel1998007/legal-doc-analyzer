"""
Text Extraction Service for Legal Documents
Handles PDF and DOCX file text extraction
"""
import os
import re
from typing import Optional
from fastapi import HTTPException, status

class ExtractionService:
    """Service for extracting text from documents"""

    @staticmethod
    def extract_text(file_path: str, file_type: str) -> str:
        """
        Extract text from a document based on its type
        
        Args:
            file_path: Path to the document file
            file_type: MIME type of the document
            
        Returns:
            Extracted text as string
        """
        if not os.path.exists(file_path):
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"File not found: {file_path}"
            )

        if file_type == "application/pdf":
            return ExtractionService.extract_from_pdf(file_path)
        
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return ExtractionService.extract_from_docx(file_path)
        
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Unsupported file type: {file_type}"
            )

    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""

        # Try pdfplumber first (better for complex layouts)
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n"
                        text += page_text
            
            if text.strip():
                print(f"✅ Extracted {len(text)} characters using pdfplumber")
                return ExtractionService.clean_text(text)

        except Exception as e:
            print(f"⚠️ pdfplumber failed: {e}, trying PyPDF2...")

        # Fallback to PyPDF2
        try:
            import PyPDF2
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n"
                        text += page_text
            
            if text.strip():
                print(f"✅ Extracted {len(text)} characters using PyPDF2")
                return ExtractionService.clean_text(text)

        except Exception as e:
            print(f"❌ PyPDF2 also failed: {e}")

        # If both fail return empty string
        print("⚠️ Could not extract text from PDF")
        return ""

    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        text = ""

        try:
            from docx import Document
            doc = Document(file_path)

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

            # Extract tables
            for table in doc.tables:
                text += "\n--- Table ---\n"
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells
                    )
                    if row_text.strip():
                        text += row_text + "\n"

            print(f"✅ Extracted {len(text)} characters from DOCX")
            return ExtractionService.clean_text(text)

        except Exception as e:
            print(f"❌ Error extracting DOCX: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error extracting text from DOCX: {str(e)}"
            )

    @staticmethod
    def clean_text(text: str) -> str:
        """
        Clean and normalize extracted text
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""

        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        
        # Remove non-printable characters (except newlines and tabs)
        text = re.sub(r'[^\x20-\x7E\n\t]', ' ', text)
        
        # Normalize spaces
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Strip leading/trailing whitespace from each line
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        # Remove excessive blank lines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()

    @staticmethod
    def get_document_info(file_path: str, file_type: str) -> dict:
        """
        Get document metadata
        
        Args:
            file_path: Path to the document
            file_type: MIME type of the document
            
        Returns:
            Dictionary with document info
        """
        info = {
            "file_path": file_path,
            "file_size": os.path.getsize(file_path),
            "file_type": file_type,
            "page_count": 0,
            "word_count": 0,
            "char_count": 0
        }

        try:
            if file_type == "application/pdf":
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    info["page_count"] = len(pdf.pages)

            elif "wordprocessingml" in file_type:
                from docx import Document
                doc = Document(file_path)
                info["page_count"] = len(doc.paragraphs)

            # Extract text to count words
            text = ExtractionService.extract_text(file_path, file_type)
            info["char_count"] = len(text)
            info["word_count"] = len(text.split())

        except Exception as e:
            print(f"Error getting document info: {e}")

        return info